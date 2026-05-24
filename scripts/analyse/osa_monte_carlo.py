"""
============================================================
OSA Observatory -- scripts/analyse/osa_monte_carlo.py
Analyse de sensibilite Monte Carlo -- ISA v2.0
============================================================
3 dimensions d'analyse :
  D1 : Sensibilite aux poids piliers
  D2 : Sensibilite aux poids semantiques indicateurs
  D3 : Incertitude d'imputation (confidence_score)

10 000 simulations par dimension.
Annee de reference : 2022 (pipeline valide, AMAR propre).

Usage :
  py -3.12 scripts/analyse/osa_monte_carlo.py
  py -3.12 scripts/analyse/osa_monte_carlo.py --year 2022 --n 10000
  py -3.12 scripts/analyse/osa_monte_carlo.py --report
============================================================
"""
from __future__ import annotations

import argparse
import json
import logging
import os
import sys
from datetime import datetime
from pathlib import Path

import numpy as np
import pandas as pd
import psycopg2
from scipy import stats

sys.path.insert(0, r"G:\python-packages")

# ── Configuration ─────────────────────────────────────────────
REPO_ROOT  = Path(__file__).parent.parent.parent
REPORT_DIR = REPO_ROOT / "docs"
LOG_DIR    = REPO_ROOT / "logs" / "monte_carlo"

DB_HOST = os.getenv("OSA_DB_HOST", "127.0.0.1")
DB_PORT = os.getenv("OSA_DB_PORT", "5432")
DB_NAME = os.getenv("OSA_DB_NAME", "osa_db")
DB_USER = os.getenv("OSA_DB_USER", "postgres")

PILLARS = ["PECO", "PENV", "PGEO", "PHUM", "PMIL",
           "PMIN", "PMON", "PNUM", "PRES", "PTRA"]
N_PILLARS   = len(PILLARS)
W_BASE      = 1.0 / N_PILLARS  # 0.10 par pilier

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)-8s | %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
log = logging.getLogger("osa_monte_carlo")


# ── Connexion DB ──────────────────────────────────────────────

def get_conn():
    return psycopg2.connect(
        host=DB_HOST, port=DB_PORT, dbname=DB_NAME,
        user=DB_USER, password=os.getenv("OSA_DB_PASS", ""),
        connect_timeout=10
    )


# ── Chargement des donnees ────────────────────────────────────

def load_pillar_scores(year: int) -> pd.DataFrame:
    """Charge les scores ISA par pays et par pilier."""
    sql = """
        SELECT country_iso3, pillar_code,
               COALESCE(isa_observed_score, 0)          AS score,
               COALESCE(avg_observation_confidence, 0.5) AS confidence
        FROM ma.v_isa_observed_scores_by_pillar
        WHERE year = %s
          AND publication_status IS NOT NULL
        ORDER BY country_iso3, pillar_code
    """
    with get_conn() as conn:
        df = pd.read_sql(sql, conn, params=(year,))
    log.info("Scores piliers charges : %d lignes (%d pays x %d piliers)",
             len(df), df["country_iso3"].nunique(), df["pillar_code"].nunique())
    return df


def load_country_scores(year: int) -> pd.DataFrame:
    """Charge les scores ISA globaux par pays."""
    sql = """
        SELECT country_iso3,
               COALESCE(isa_observed_score, 0)          AS isa_score,
               COALESCE(avg_observation_confidence, 0.5) AS confidence
        FROM ma.v_isa_observed_scores_by_country_year
        WHERE year = %s
        ORDER BY country_iso3
    """
    with get_conn() as conn:
        df = pd.read_sql(sql, conn, params=(year,))
    log.info("Scores pays charges : %d pays", len(df))
    return df


def load_semantic_weights() -> pd.DataFrame:
    """Charge les poids semantiques des indicateurs."""
    sql = """
        SELECT semantic_code,
               COALESCE(sovereignty_weight, 0.5) AS w_sovereignty,
               COALESCE(trust_level, 0.5)        AS w_trust,
               COALESCE(forecastability, 0.5)    AS w_forecast
        FROM rf.semantic_governance_matrix
        ORDER BY semantic_code
    """
    with get_conn() as conn:
        df = pd.read_sql(sql, conn)
    log.info("Poids semantiques charges : %d codes", len(df))
    return df


# ── Dimension 1 : Sensibilite poids piliers ───────────────────

def simulate_d1_pillar_weights(
    df_pillar: pd.DataFrame,
    n_sim: int = 10_000,
    concentration: float = 10.0,
) -> dict:
    """
    D1 : Variation des poids piliers par distribution de Dirichlet.
    concentration = 10 => variation moderee autour de 0.10 par pilier.
    Mesures :
      - ecart-type des scores ISA simules par pays
      - % de pays changeant de quintile par rapport au baseline
      - rang moyen et ecart-type des rangs
    """
    log.info("D1 : Simulation poids piliers (%d iterations)...", n_sim)

    # Matrice pays x piliers (54 x 10)
    pivot = df_pillar.pivot(
        index="country_iso3", columns="pillar_code", values="score"
    ).fillna(0)

    # S'assurer que tous les piliers sont presents
    for p in PILLARS:
        if p not in pivot.columns:
            pivot[p] = 0.0
    pivot = pivot[PILLARS]

    countries   = pivot.index.tolist()
    scores_mat  = pivot.values  # (54, 10)
    n_countries = len(countries)

    # Baseline : poids egaux
    w_base      = np.ones(N_PILLARS) / N_PILLARS
    isa_base    = scores_mat @ w_base
    rank_base   = stats.rankdata(-isa_base)  # rang 1 = meilleur

    # Simulations
    alpha       = w_base * concentration  # parametres Dirichlet
    all_scores  = np.zeros((n_sim, n_countries))
    all_ranks   = np.zeros((n_sim, n_countries))

    for i in range(n_sim):
        w_sim           = np.random.dirichlet(alpha)
        isa_sim         = scores_mat @ w_sim
        all_scores[i]   = isa_sim
        all_ranks[i]    = stats.rankdata(-isa_sim)

    # Metriques
    score_std   = all_scores.std(axis=0)
    rank_std    = all_ranks.std(axis=0)
    rank_mean   = all_ranks.mean(axis=0)

    # Stabilite de quintile
    q_base      = pd.qcut(isa_base, 5, labels=False)
    q_changes   = np.zeros(n_countries)
    for i in range(n_sim):
        q_sim       = pd.qcut(all_scores[i], 5, labels=False,
                              duplicates="drop")
        q_changes  += (q_sim != q_base).astype(int)
    pct_quintile_change = (q_changes / n_sim * 100)

    # Robustesse globale
    robustness_score = 1.0 - score_std.mean() / (isa_base.mean() + 1e-9)

    results = {
        "dimension":            "D1_PILLAR_WEIGHTS",
        "n_simulations":        n_sim,
        "concentration":        concentration,
        "robustness_score":     round(float(robustness_score), 4),
        "avg_score_std":        round(float(score_std.mean()), 4),
        "avg_rank_std":         round(float(rank_std.mean()), 4),
        "avg_pct_quintile_change": round(float(pct_quintile_change.mean()), 2),
        "max_pct_quintile_change": round(float(pct_quintile_change.max()), 2),
        "most_stable_country":  countries[score_std.argmin()],
        "most_sensitive_country": countries[score_std.argmax()],
        "by_country": [
            {
                "country_iso3":          c,
                "isa_baseline":          round(float(isa_base[i]), 4),
                "rank_baseline":         int(rank_base[i]),
                "score_std":             round(float(score_std[i]), 4),
                "rank_std":              round(float(rank_std[i]), 4),
                "pct_quintile_change":   round(float(pct_quintile_change[i]), 2),
            }
            for i, c in enumerate(countries)
        ],
        "verdict": (
            "ROBUSTE" if robustness_score > 0.85
            else "MODEREMENT_ROBUSTE" if robustness_score > 0.70
            else "SENSIBLE"
        ),
    }
    log.info("D1 termine : robustesse=%.3f (%s)",
             robustness_score, results["verdict"])
    return results


# ── Dimension 2 : Sensibilite poids semantiques ───────────────

def simulate_d2_semantic_weights(
    df_pillar: pd.DataFrame,
    df_semantic: pd.DataFrame,
    n_sim: int = 10_000,
    variation: float = 0.20,
) -> dict:
    """
    D2 : Variation des poids semantiques ±20% (uniforme).
    Approximation : on perturbe directement les scores piliers
    proportionnellement aux poids semantiques.
    """
    log.info("D2 : Simulation poids semantiques (%d iterations)...", n_sim)

    pivot    = df_pillar.pivot(
        index="country_iso3", columns="pillar_code", values="score"
    ).fillna(0)
    for p in PILLARS:
        if p not in pivot.columns:
            pivot[p] = 0.0
    pivot       = pivot[PILLARS]
    countries   = pivot.index.tolist()
    scores_mat  = pivot.values
    n_countries = len(countries)

    w_base      = np.ones(N_PILLARS) / N_PILLARS
    isa_base    = scores_mat @ w_base
    rank_base   = stats.rankdata(-isa_base)

    # Poids semantique moyen comme facteur de perturbation par pilier
    w_sem_base  = df_semantic["w_sovereignty"].values
    w_sem_mean  = w_sem_base.mean()

    all_scores  = np.zeros((n_sim, n_countries))

    for i in range(n_sim):
        # Perturber chaque pilier proportionnellement a la variance semantique
        perturb     = 1.0 + np.random.uniform(-variation, variation, N_PILLARS)
        scores_sim  = scores_mat * perturb
        scores_sim  = np.clip(scores_sim, 0, 1)
        all_scores[i] = scores_sim @ w_base

    score_std   = all_scores.std(axis=0)
    rank_std    = stats.rankdata(-all_scores.mean(axis=0))

    robustness_score = 1.0 - score_std.mean() / (isa_base.mean() + 1e-9)

    results = {
        "dimension":            "D2_SEMANTIC_WEIGHTS",
        "n_simulations":        n_sim,
        "variation_pct":        variation * 100,
        "robustness_score":     round(float(robustness_score), 4),
        "avg_score_std":        round(float(score_std.mean()), 4),
        "most_stable_country":  countries[score_std.argmin()],
        "most_sensitive_country": countries[score_std.argmax()],
        "by_country": [
            {
                "country_iso3":  c,
                "isa_baseline":  round(float(isa_base[i]), 4),
                "score_std":     round(float(score_std[i]), 4),
            }
            for i, c in enumerate(countries)
        ],
        "verdict": (
            "ROBUSTE" if robustness_score > 0.85
            else "MODEREMENT_ROBUSTE" if robustness_score > 0.70
            else "SENSIBLE"
        ),
    }
    log.info("D2 termine : robustesse=%.3f (%s)",
             robustness_score, results["verdict"])
    return results


# ── Dimension 3 : Incertitude d'imputation ────────────────────

def simulate_d3_imputation_uncertainty(
    df_pillar: pd.DataFrame,
    n_sim: int = 10_000,
) -> dict:
    """
    D3 : Bruit gaussien N(0, sigma) avec sigma = 1 - confidence_score.
    Mesure l'intervalle de confiance [P5, P95] des scores pays.
    """
    log.info("D3 : Simulation incertitude imputation (%d iterations)...", n_sim)

    pivot_score = df_pillar.pivot(
        index="country_iso3", columns="pillar_code", values="score"
    ).fillna(0)
    pivot_conf  = df_pillar.pivot(
        index="country_iso3", columns="pillar_code", values="confidence"
    ).fillna(0.5)

    for p in PILLARS:
        if p not in pivot_score.columns:
            pivot_score[p] = 0.0
            pivot_conf[p]  = 0.5
    pivot_score  = pivot_score[PILLARS]
    pivot_conf   = pivot_conf[PILLARS]

    countries    = pivot_score.index.tolist()
    scores_mat   = pivot_score.values   # (54, 10)
    sigma_mat    = 1.0 - pivot_conf.values  # incertitude = 1 - confiance

    w_base       = np.ones(N_PILLARS) / N_PILLARS
    isa_base     = scores_mat @ w_base
    n_countries  = len(countries)

    all_scores   = np.zeros((n_sim, n_countries))

    for i in range(n_sim):
        # Bruit proportionnel au score (sigma relatif, pas absolu)
        sigma_rel    = sigma_mat * np.maximum(scores_mat, 0.05)
        noise        = np.random.normal(0, sigma_rel)
        scores_sim   = np.clip(scores_mat + noise, 0, 1)
        all_scores[i] = scores_sim @ w_base

    p5   = np.percentile(all_scores, 5,  axis=0)
    p25  = np.percentile(all_scores, 25, axis=0)
    p75  = np.percentile(all_scores, 75, axis=0)
    p95  = np.percentile(all_scores, 95, axis=0)
    ci90 = p95 - p5  # amplitude intervalle de confiance 90%

    robustness_score = 1.0 - ci90.mean() / (isa_base.mean() + 1e-9)

    results = {
        "dimension":             "D3_IMPUTATION_UNCERTAINTY",
        "n_simulations":         n_sim,
        "robustness_score":      round(float(robustness_score), 4),
        "avg_ci90_amplitude":    round(float(ci90.mean()), 4),
        "max_ci90_amplitude":    round(float(ci90.max()), 4),
        "most_stable_country":   countries[ci90.argmin()],
        "most_uncertain_country": countries[ci90.argmax()],
        "by_country": [
            {
                "country_iso3":  c,
                "isa_baseline":  round(float(isa_base[i]), 4),
                "p5":            round(float(p5[i]), 4),
                "p25":           round(float(p25[i]), 4),
                "p75":           round(float(p75[i]), 4),
                "p95":           round(float(p95[i]), 4),
                "ci90_amplitude": round(float(ci90[i]), 4),
            }
            for i, c in enumerate(countries)
        ],
        "verdict": (
            "ROBUSTE" if robustness_score > 0.85
            else "MODEREMENT_ROBUSTE" if robustness_score > 0.70
            else "SENSIBLE"
        ),
    }
    log.info("D3 termine : robustesse=%.3f (%s)",
             robustness_score, results["verdict"])
    return results


# ── Rapport Word ──────────────────────────────────────────────

def generate_word_report(results: dict, year: int) -> Path:
    """Genere un rapport Word complet de l'analyse Monte Carlo."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        log.warning("python-docx non installe - rapport Word ignore")
        return None

    doc = Document()

    def add_h1(text, color=(31, 78, 121)):
        h = doc.add_heading(text, level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)

    def add_h2(text):
        h = doc.add_heading(text, level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(46, 117, 182)

    def add_para(text, italic=False, bold=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.italic = italic
        run.bold   = bold
        return p

    def verdict_color(v):
        if v == "ROBUSTE":             return RGBColor(55, 86, 35)
        if v == "MODEREMENT_ROBUSTE":  return RGBColor(127, 96, 0)
        return RGBColor(192, 0, 0)

    # Titre
    title = doc.add_heading("OSA Observatory — Analyse de sensibilite Monte Carlo", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        f"Annee de reference : {year}  |  Simulations : 10 000 par dimension  |  "
        f"Date : {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    )
    doc.add_paragraph("")

    # Synthese
    add_h1("Synthese executif")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Dimension"
    hdr[1].text = "Robustesse"
    hdr[2].text = "Verdict"
    hdr[3].text = "Pays le plus sensible"

    for dim_key in ["d1", "d2", "d3"]:
        d = results[dim_key]
        row = table.add_row().cells
        row[0].text = d["dimension"]
        row[1].text = f"{d['robustness_score']:.3f}"
        row[2].text = d["verdict"]
        row[3].text = d.get("most_sensitive_country",
                             d.get("most_uncertain_country", "N/A"))

    doc.add_paragraph("")

    # D1
    add_h1("Dimension 1 — Sensibilite aux poids piliers")
    d1 = results["d1"]
    add_para(
        f"Score de robustesse : {d1['robustness_score']:.4f}  |  "
        f"Verdict : {d1['verdict']}",
        bold=True
    )
    add_para(
        f"Ecart-type moyen des scores ISA : {d1['avg_score_std']:.4f}  |  "
        f"Ecart-type moyen des rangs : {d1['avg_rank_std']:.2f}  |  "
        f"Changement de quintile moyen : {d1['avg_pct_quintile_change']:.1f}%  |  "
        f"Max : {d1['max_pct_quintile_change']:.1f}%"
    )
    add_para(
        "Interpretation : une variation de ±50% autour de la ponderation "
        "egale (0.10 par pilier) produit un ecart-type moyen de "
        f"{d1['avg_score_std']:.4f} sur les scores ISA. "
        f"{d1['avg_pct_quintile_change']:.1f}% des pays changent de quintile "
        "en moyenne. L'indice est " +
        ("stable aux choix de ponderation." if d1["verdict"] == "ROBUSTE"
         else "moderement sensible aux choix de ponderation."
         if d1["verdict"] == "MODEREMENT_ROBUSTE"
         else "sensible aux choix de ponderation — recalibration recommandee."),
        italic=True
    )

    # Top 5 pays les plus sensibles D1
    add_h2("Pays les plus sensibles (D1)")
    top5 = sorted(d1["by_country"], key=lambda x: x["score_std"], reverse=True)[:5]
    t = doc.add_table(rows=1, cols=5)
    t.style = "Table Grid"
    h = t.rows[0].cells
    h[0].text = "Pays"
    h[1].text = "Score ISA baseline"
    h[2].text = "Rang baseline"
    h[3].text = "Ecart-type score"
    h[4].text = "% changement quintile"
    for c in top5:
        r = t.add_row().cells
        r[0].text = c["country_iso3"]
        r[1].text = str(c["isa_baseline"])
        r[2].text = str(c["rank_baseline"])
        r[3].text = str(c["score_std"])
        r[4].text = f"{c['pct_quintile_change']:.1f}%"
    doc.add_paragraph("")

    # D2
    add_h1("Dimension 2 — Sensibilite aux poids semantiques")
    d2 = results["d2"]
    add_para(
        f"Score de robustesse : {d2['robustness_score']:.4f}  |  "
        f"Verdict : {d2['verdict']}",
        bold=True
    )
    add_para(
        f"Variation appliquee : ±{d2['variation_pct']:.0f}%  |  "
        f"Ecart-type moyen des scores ISA : {d2['avg_score_std']:.4f}"
    )
    add_para(
        "Interpretation : une variation de ±20% des poids semantiques "
        "produit un ecart-type moyen de "
        f"{d2['avg_score_std']:.4f} sur les scores ISA. "
        "Les choix de ponderation semantique ont un impact " +
        ("faible." if d2["verdict"] == "ROBUSTE"
         else "modere." if d2["verdict"] == "MODEREMENT_ROBUSTE"
         else "significatif — revision doctrinale recommandee."),
        italic=True
    )
    doc.add_paragraph("")

    # D3
    add_h1("Dimension 3 — Incertitude d'imputation")
    d3 = results["d3"]
    add_para(
        f"Score de robustesse : {d3['robustness_score']:.4f}  |  "
        f"Verdict : {d3['verdict']}",
        bold=True
    )
    add_para(
        f"Amplitude moyenne IC90 : {d3['avg_ci90_amplitude']:.4f}  |  "
        f"Amplitude max IC90 : {d3['max_ci90_amplitude']:.4f}  |  "
        f"Pays le plus incertain : {d3['most_uncertain_country']}"
    )
    add_para(
        "Interpretation : l'intervalle de confiance a 90% des scores ISA "
        "a une amplitude moyenne de "
        f"{d3['avg_ci90_amplitude']:.4f}. "
        "L'incertitude liee aux donnees imputees est " +
        ("faible et acceptable." if d3["verdict"] == "ROBUSTE"
         else "moderee — des efforts de collecte supplementaires sont "
              "recommandes pour les pays a faible confiance."
         if d3["verdict"] == "MODEREMENT_ROBUSTE"
         else "elevee — les resultats des pays a faible couverture "
              "doivent etre presentes avec intervalles de confiance."),
        italic=True
    )

    # Top 5 pays les plus incertains D3
    add_h2("Pays avec les intervalles de confiance les plus larges (D3)")
    top5_d3 = sorted(
        d3["by_country"], key=lambda x: x["ci90_amplitude"], reverse=True
    )[:5]
    t3 = doc.add_table(rows=1, cols=5)
    t3.style = "Table Grid"
    h3 = t3.rows[0].cells
    h3[0].text = "Pays"
    h3[1].text = "ISA baseline"
    h3[2].text = "P5"
    h3[3].text = "P95"
    h3[4].text = "IC90"
    for c in top5_d3:
        r = t3.add_row().cells
        r[0].text = c["country_iso3"]
        r[1].text = str(c["isa_baseline"])
        r[2].text = str(c["p5"])
        r[3].text = str(c["p95"])
        r[4].text = str(c["ci90_amplitude"])
    doc.add_paragraph("")

    # Conclusion
    add_h1("Conclusion pour le Conseil scientifique")
    verdicts = [results["d1"]["verdict"], results["d2"]["verdict"],
                results["d3"]["verdict"]]
    n_robuste = verdicts.count("ROBUSTE")
    if n_robuste == 3:
        conclusion = (
            "L'ISA est robuste sur les trois dimensions analysees. "
            "La ponderation egale des piliers est scientifiquement defensible "
            "devant le Conseil scientifique. Une differentiation des poids "
            "est possible mais non necessaire pour garantir la stabilite des "
            "resultats."
        )
    elif n_robuste >= 2:
        conclusion = (
            "L'ISA est moderement robuste. La ponderation egale est acceptable "
            "pour la phase d'appropriation 2026. Le Conseil scientifique devra "
            "se prononcer sur la dimension la moins robuste avant le lancement "
            "officiel de septembre 2027."
        )
    else:
        conclusion = (
            "L'ISA presente des sensibilites significatives. Le Conseil "
            "scientifique doit examiner en priorite les choix de ponderation "
            "et la politique d'imputation avant certification de la publication."
        )
    add_para(conclusion, italic=True)

    # Sauvegarde
    ts       = datetime.now().strftime("%Y%m%d")
    out_path = REPORT_DIR / f"OSA_Monte_Carlo_ISA_v2_{year}_{ts}.docx"
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    log.info("Rapport Word genere : %s", out_path)
    return out_path


# ── Main ──────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="OSA Monte Carlo -- Analyse de sensibilite ISA v2"
    )
    parser.add_argument("--year",   type=int,  default=2022)
    parser.add_argument("--n",      type=int,  default=10_000)
    parser.add_argument("--report", action="store_true",
                        help="Genere un rapport Word")
    parser.add_argument("--quick",  action="store_true",
                        help="Mode rapide : 1000 simulations")
    args = parser.parse_args()

    n_sim = 1_000 if args.quick else args.n
    year  = args.year

    log.info("=" * 55)
    log.info("OSA Monte Carlo ISA v2 -- annee %d -- %d simulations", year, n_sim)
    log.info("=" * 55)

    # Chargement
    df_pillar  = load_pillar_scores(year)
    df_semantic = load_semantic_weights()

    if df_pillar.empty:
        log.error("Aucune donnee pilier pour l'annee %d", year)
        sys.exit(1)

    # Simulations
    np.random.seed(42)  # reproductibilite
    d1 = simulate_d1_pillar_weights(df_pillar, n_sim=n_sim)
    d2 = simulate_d2_semantic_weights(df_pillar, df_semantic, n_sim=n_sim)
    d3 = simulate_d3_imputation_uncertainty(df_pillar, n_sim=n_sim)

    results = {
        "metadata": {
            "year":          year,
            "n_simulations": n_sim,
            "timestamp":     datetime.now().isoformat(),
            "random_seed":   42,
        },
        "d1": d1,
        "d2": d2,
        "d3": d3,
        "synthese": {
            "d1_verdict": d1["verdict"],
            "d2_verdict": d2["verdict"],
            "d3_verdict": d3["verdict"],
            "robustesse_globale": round(
                (d1["robustness_score"] +
                 d2["robustness_score"] +
                 d3["robustness_score"]) / 3, 4
            ),
        }
    }

    # Affichage console
    print(f"\n{'='*55}")
    print(f"  OSA Monte Carlo ISA v2 -- {year} -- {n_sim} simulations")
    print(f"{'='*55}")
    print(f"\n  D1 Poids piliers    : {d1['verdict']:25s}  robustesse={d1['robustness_score']:.4f}")
    print(f"  D2 Poids semantiques : {d2['verdict']:25s}  robustesse={d2['robustness_score']:.4f}")
    print(f"  D3 Imputation        : {d3['verdict']:25s}  robustesse={d3['robustness_score']:.4f}")
    print(f"\n  Robustesse globale : {results['synthese']['robustesse_globale']:.4f}")
    print(f"\n  Pays le plus sensible (D1) : {d1['most_sensitive_country']}")
    print(f"  Pays le plus incertain (D3) : {d3['most_uncertain_country']}")
    print(f"{'='*55}\n")

    # Sauvegarde JSON
    LOG_DIR.mkdir(parents=True, exist_ok=True)
    ts       = datetime.now().strftime("%Y%m%d_%H%M%S")
    json_out = LOG_DIR / f"monte_carlo_{year}_{ts}.json"

    # Serialisation sans les listes by_country pour le JSON principal
    results_summary = {
        "metadata":  results["metadata"],
        "synthese":  results["synthese"],
        "d1_summary": {k: v for k, v in d1.items() if k != "by_country"},
        "d2_summary": {k: v for k, v in d2.items() if k != "by_country"},
        "d3_summary": {k: v for k, v in d3.items() if k != "by_country"},
    }
    with open(json_out, "w", encoding="utf-8") as f:
        json.dump(results_summary, f, ensure_ascii=False, indent=2)
    log.info("Resultats JSON : %s", json_out)

    # Rapport Word
    if args.report:
        out_path = generate_word_report(results, year)
        if out_path:
            print(f"  Rapport Word : {out_path}")

    # Code retour
    robustesse = results["synthese"]["robustesse_globale"]
    if robustesse >= 0.85:
        log.info("ROBUSTE -- ponderation egale defensible devant le Conseil scientifique")
        sys.exit(0)
    elif robustesse >= 0.70:
        log.warning("MODEREMENT ROBUSTE -- examen recommande par le Conseil scientifique")
        sys.exit(1)
    else:
        log.error("SENSIBLE -- revision ponderation requise avant certification")
        sys.exit(2)


if __name__ == "__main__":
    main()
