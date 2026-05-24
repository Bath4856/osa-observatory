"""
============================================================
OSA Observatory -- scripts/analyse/osa_healthcheck.py
Analyse automatique du healthcheck pipeline OSA
============================================================
Usage :
  py -3.12 scripts/analyse/osa_healthcheck.py
  py -3.12 scripts/analyse/osa_healthcheck.py --log logs/healthcheck/healthcheck_20260523.log
  py -3.12 scripts/analyse/osa_healthcheck.py --report  # genere rapport Word si anomalies
============================================================
"""
from __future__ import annotations

import sys
sys.path.insert(0, r"G:\python-packages")
import argparse
import json
import logging
import os
import subprocess
import sys
from datetime import datetime
from pathlib import Path


# ── Configuration ─────────────────────────────────────────────
REPO_ROOT   = Path(__file__).parent.parent.parent
LOG_DIR     = REPO_ROOT / "logs" / "healthcheck"
AUDIT_SQL   = REPO_ROOT / "audit" / "osa_healthcheck.sql"
REPORT_DIR  = REPO_ROOT / "docs"

DB_HOST = os.getenv("OSA_DB_HOST", "127.0.0.1")
DB_PORT = os.getenv("OSA_DB_PORT", "5432")
DB_NAME = os.getenv("OSA_DB_NAME", "osa_db")
DB_USER = os.getenv("OSA_DB_USER", "postgres")

# ── Seuils d'alerte ───────────────────────────────────────────
SEUILS = {
    "doctrine_compliance_max":    0,      # zero tolerance
    "active_no_l3_max":           0,      # zero tolerance
    "duplicates_l1_max":          0,      # zero tolerance
    "source_null_max":            0,      # zero tolerance
    "amar_break_min":             3,      # saut min pays pour alerte
    "isa_score_variation_max":    0.05,   # variation max avg ISA entre annees
    "green_annees_recentes_max":  0,      # 0 GREEN attendu apres 2019
}

# ── Logging ───────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)-8s | %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
log = logging.getLogger("osa_healthcheck")


# ── Execution SQL ─────────────────────────────────────────────

def run_healthcheck_sql(log_file: Path) -> list[str]:
    """Execute le SQL de healthcheck et retourne les lignes de resultat."""
    log_file.parent.mkdir(parents=True, exist_ok=True)
    log.info("Execution SQL healthcheck -> %s", log_file)

    result = subprocess.run(
        ["psql", "-h", DB_HOST, "-p", DB_PORT, "-U", DB_USER, "-d", DB_NAME,
         "-f", str(AUDIT_SQL)],
        capture_output=True, text=True, encoding="utf-8"
    )

    if result.returncode != 0:
        log.error("Erreur psql : %s", result.stderr)
        raise RuntimeError(f"psql failed : {result.stderr}")

    output = result.stdout
    with open(log_file, "w", encoding="utf-8") as f:
        f.write(f"# OSA Healthcheck -- {datetime.now().isoformat()}\n")
        f.write(f"# DB : {DB_NAME}@{DB_HOST}:{DB_PORT}\n\n")
        f.write(output)

    lines = [l.strip() for l in output.splitlines() if "|" in l and not l.startswith("#")]
    log.info("%d lignes de diagnostic recuperees", len(lines))
    return lines


# ── Parsers par type de check ─────────────────────────────────

def parse_lines(lines: list[str]) -> dict:
    """Parse les lignes de resultat en structure de donnees."""
    data = {
        "views":         {},
        "doctrine":      {"count": 0, "codes": []},
        "active_no_l3":  {"count": 0, "codes": []},
        "duplicates":    0,
        "source_null":   0,
        "isa_scores":    {},
        "amar_dist":     {},
        "geneco_dist":   {},
        "amar_breaks":   [],
        "piliers":       {},
        "layers":        {},
        "versions":      [],
        "known_breaks":  [],
    }

    for line in lines:
        parts = [p.strip() for p in line.split("|")]
        if not parts:
            continue
        key = parts[0]

        if key == "CHECK_VIEWS":
            labels = ["p7i_source","p7i_core","p7i_escalation",
                      "amar_dashboard","geneco_dashboard","composite_dashboard",
                      "amar_public","geneco_public"]
            for i, label in enumerate(labels):
                data["views"][label] = parts[i+1] if i+1 < len(parts) else "UNKNOWN"

        elif key == "CHECK_DOCTRINE":
            data["doctrine"]["count"] = int(parts[1]) if parts[1].isdigit() else 0
            data["doctrine"]["codes"] = parts[2].split(",") if len(parts) > 2 and parts[2] else []

        elif key == "CHECK_ACTIVE_NO_L3":
            data["active_no_l3"]["count"] = int(parts[1]) if parts[1].isdigit() else 0
            data["active_no_l3"]["codes"] = parts[2].split(",") if len(parts) > 2 and parts[2] else []

        elif key == "CHECK_DUPLICATES_L1":
            data["duplicates"] = int(parts[1]) if len(parts) > 1 and parts[1].isdigit() else 0

        elif key == "CHECK_SOURCE_NULL":
            data["source_null"] = int(parts[1]) if len(parts) > 1 and parts[1].isdigit() else 0

        elif key == "ISA_SCORES" and len(parts) >= 4:
            year = int(parts[1])
            data["isa_scores"][year] = {
                "avg": float(parts[2]) if parts[2] else 0.0,
                "nb":  int(parts[3]) if parts[3].isdigit() else 0,
            }

        elif key == "AMAR_DIST" and len(parts) >= 5:
            year = int(parts[1])
            band = parts[2]
            if year not in data["amar_dist"]:
                data["amar_dist"][year] = {}
            data["amar_dist"][year][band] = {
                "nb":    int(parts[3]) if parts[3].isdigit() else 0,
                "score": float(parts[4]) if parts[4] else 0.0,
            }

        elif key == "GENECO_DIST" and len(parts) >= 5:
            year = int(parts[1])
            band = parts[2]
            if year not in data["geneco_dist"]:
                data["geneco_dist"][year] = {}
            data["geneco_dist"][year][band] = {
                "nb":    int(parts[3]) if parts[3].isdigit() else 0,
                "score": float(parts[4]) if parts[4] else 0.0,
            }

        elif key == "AMAR_BREAK" and len(parts) >= 6:
            data["amar_breaks"].append({
                "year":      int(parts[1]),
                "band":      parts[2],
                "nb_pays":   int(parts[3]) if parts[3].isdigit() else 0,
                "prev_nb":   int(parts[4]) if parts[4].isdigit() else 0,
                "delta":     int(parts[5]) if parts[5].isdigit() else 0,
            })

        elif key == "PILIER_COUNT" and len(parts) >= 3:
            data["piliers"][parts[1]] = int(parts[2]) if parts[2].isdigit() else 0

        elif key == "LAYER_COUNT" and len(parts) >= 4:
            lid = int(parts[1])
            data["layers"][lid] = {
                "lignes":     int(parts[2]) if parts[2].isdigit() else 0,
                "indicateurs":int(parts[3]) if parts[3].isdigit() else 0,
                "pays":       int(parts[4]) if len(parts) > 4 and parts[4].isdigit() else 0,
            }

        elif key == "AMAR_KNOWN_BREAKS" and len(parts) >= 7:
            data["known_breaks"].append({
                "id":        int(parts[1]) if parts[1].isdigit() else 0,
                "year_from": int(parts[2]) if parts[2].isdigit() else 0,
                "year_to":   int(parts[3]) if parts[3].isdigit() else 0,
                "band":      parts[4],
                "cause":     parts[5],
                "artefact":  parts[6].lower() == "true",
            })

        elif key == "VERSIONS_COUNT" and len(parts) >= 4:
            data["versions"].append({
                "sprint": parts[1],
                "action": parts[2],
                "count":  int(parts[3]) if parts[3].isdigit() else 0,
            })

    return data


# ── Analyse et detection d'anomalies ─────────────────────────

def analyse(data: dict) -> dict:
    """Analyse les donnees et produit un rapport d'anomalies."""
    anomalies   = []
    avertissements = []
    ok          = []

    # ── Vues critiques
    manquantes = [k for k, v in data["views"].items() if v == "MISSING"]
    if manquantes:
        anomalies.append({
            "code": "VIEWS_MISSING",
            "niveau": "CRITIQUE",
            "message": f"Vues critiques manquantes : {', '.join(manquantes)}",
            "action": "Relancer restart_v2.ps1 option [7] pour redeployer AMAR + GENECO"
        })
    else:
        ok.append("Toutes les vues critiques sont presentes")

    # ── Conformite doctrinale
    if data["doctrine"]["count"] > SEUILS["doctrine_compliance_max"]:
        anomalies.append({
            "code": "DOCTRINE_NON_CONFORME",
            "niveau": "CRITIQUE",
            "message": f"{data['doctrine']['count']} indicateur(s) actif(s) non conformes Doctrine ISA v1 : {', '.join(data['doctrine']['codes'])}",
            "action": "Desactiver ou remplacer les indicateurs non conformes via rf.indicator_versions"
        })
    else:
        ok.append("Conformite doctrinale 100% - zero indicateur de perception en production")

    # ── Indicateurs actifs sans L3
    if data["active_no_l3"]["count"] > SEUILS["active_no_l3_max"]:
        avertissements.append({
            "code": "ACTIVE_NO_L3",
            "niveau": "ATTENTION",
            "message": f"{data['active_no_l3']['count']} indicateur(s) actif(s) sans donnees L3 : {', '.join(data['active_no_l3']['codes'])}",
            "action": "Verifier normalize_indicator() ou desactiver l'indicateur"
        })
    else:
        ok.append("Tous les indicateurs actifs ont des donnees L3")

    # ── Doublons L1
    if data["duplicates"] > SEUILS["duplicates_l1_max"]:
        anomalies.append({
            "code": "DUPLICATES_L1",
            "niveau": "CRITIQUE",
            "message": f"{data['duplicates']} doublon(s) detecte(s) en L1",
            "action": "Executer patch_deduplicate_l1_fix_constraint.sql"
        })
    else:
        ok.append("Zero doublon en L1")

    # ── Source NULL en L1
    if data["source_null"] > SEUILS["source_null_max"]:
        anomalies.append({
            "code": "SOURCE_NULL_L1",
            "niveau": "CRITIQUE",
            "message": f"{data['source_null']} ligne(s) L1 sans source_id",
            "action": "Executer patch_fix_source_id_null.sql"
        })
    else:
        ok.append("Toutes les lignes L1 ont un source_id trace")

    # ── Variation ISA entre annees consecutives
    annees = sorted(data["isa_scores"].keys())
    for i in range(1, len(annees)):
        a1, a2 = annees[i-1], annees[i]
        if a1 not in data["isa_scores"] or a2 not in data["isa_scores"]:
            continue
        delta = abs(data["isa_scores"][a2]["avg"] - data["isa_scores"][a1]["avg"])
        if delta > SEUILS["isa_score_variation_max"]:
            avertissements.append({
                "code": "ISA_SCORE_VARIATION",
                "niveau": "ATTENTION",
                "message": f"Variation ISA moyenne {a1}->{a2} : {delta:.4f} (seuil {SEUILS['isa_score_variation_max']})",
                "action": "Verifier recalcul L3 - possible artefact de recalibration"
            })

    # ── Ruptures AMAR (hors ruptures connues et documentees)
    known = {(k["year_to"], k["band"]) for k in data["known_breaks"]}
    for b in data["amar_breaks"]:
        if b["delta"] >= SEUILS["amar_break_min"]:
            if (b["year"], b["band"]) in known:
                ok.append(f"Rupture AMAR {b['band']} {b['year']-1}->{b['year']} documentee et justifiee (rf.amar_known_breaks)")
            else:
                avertissements.append({
                    "code": "AMAR_BREAK",
                    "niveau": "ATTENTION",
                    "message": f"Rupture AMAR {b['band']} en {b['year']} : {b['prev_nb']} -> {b['nb_pays']} pays (delta={b['delta']})",
                    "action": "Verifier recalibration seuils AMAR - possible artefact post-desactivation indicateurs"
                })

    # ── GREEN sur annees recentes (2020+)
    for year in range(2020, 2025):
        if year in data["amar_dist"]:
            nb_green = data["amar_dist"][year].get("GREEN", {}).get("nb", 0)
            if nb_green > SEUILS["green_annees_recentes_max"]:
                avertissements.append({
                    "code": "AMAR_GREEN_RECENT",
                    "niveau": "ATTENTION",
                    "message": f"{nb_green} pays GREEN en {year} - inattendu apres recalibration Sprint 10/11",
                    "action": "Verifier seuils AMAR - recalibration prevue Sprint 12"
                })

    # ── Summary
    return {
        "timestamp":       datetime.now().isoformat(),
        "anomalies":       anomalies,
        "avertissements":  avertissements,
        "ok":              ok,
        "statut_global":   "CRITIQUE" if anomalies else ("ATTENTION" if avertissements else "OK"),
        "data":            data,
    }


# ── Rapport texte ─────────────────────────────────────────────

def print_rapport(rapport: dict) -> None:
    """Affiche le rapport dans la console."""
    ts = rapport["timestamp"]
    statut = rapport["statut_global"]

    couleurs = {"CRITIQUE": "\033[91m", "ATTENTION": "\033[93m", "OK": "\033[92m"}
    reset = "\033[0m"
    c = couleurs.get(statut, "")

    print(f"\n{'='*60}")
    print(f"  OSA Observatory -- Healthcheck -- {ts[:19]}")
    print(f"  Statut global : {c}{statut}{reset}")
    print(f"{'='*60}")

    if rapport["anomalies"]:
        print(f"\n  {couleurs['CRITIQUE']}ANOMALIES CRITIQUES ({len(rapport['anomalies'])}){reset}")
        for a in rapport["anomalies"]:
            print(f"\n  [{a['code']}]")
            print(f"  Message : {a['message']}")
            print(f"  Action  : {a['action']}")

    if rapport["avertissements"]:
        print(f"\n  {couleurs['ATTENTION']}AVERTISSEMENTS ({len(rapport['avertissements'])}){reset}")
        for a in rapport["avertissements"]:
            print(f"\n  [{a['code']}]")
            print(f"  Message : {a['message']}")
            print(f"  Action  : {a['action']}")

    if rapport["ok"]:
        print(f"\n  {couleurs['OK']}CHECKS OK ({len(rapport['ok'])}){reset}")
        for msg in rapport["ok"]:
            print(f"  OK : {msg}")

    # Distribution AMAR 2024
    d = rapport["data"]
    if 2024 in d["amar_dist"]:
        print(f"\n  Distribution AMAR 2024 :")
        for band, vals in sorted(d["amar_dist"][2024].items()):
            print(f"    {band:8s} : {vals['nb']:3d} pays  avg={vals['score']:.3f}")

    # Scores ISA recents
    print(f"\n  Scores ISA observes (2020-2024) :")
    for y in range(2020, 2025):
        if y in d["isa_scores"]:
            s = d["isa_scores"][y]
            print(f"    {y} : avg={s['avg']:.5f}  nb={s['nb']}")

    print(f"\n{'='*60}\n")


# ── Sauvegarde JSON ───────────────────────────────────────────

def save_json(rapport: dict, log_file: Path) -> None:
    """Sauvegarde le rapport JSON a cote du log SQL."""
    json_file = log_file.with_suffix(".json")
    rapport_export = {k: v for k, v in rapport.items() if k != "data"}
    rapport_export["summary"] = {
        "nb_anomalies":      len(rapport["anomalies"]),
        "nb_avertissements": len(rapport["avertissements"]),
        "nb_ok":             len(rapport["ok"]),
        "statut_global":     rapport["statut_global"],
        "indicateurs_actifs_total": sum(rapport["data"]["piliers"].values()),
        "amar_2024": rapport["data"]["amar_dist"].get(2024, {}),
        "isa_2024":  rapport["data"]["isa_scores"].get(2024, {}),
    }
    with open(json_file, "w", encoding="utf-8") as f:
        json.dump(rapport_export, f, ensure_ascii=False, indent=2)
    log.info("Rapport JSON sauvegarde : %s", json_file)


# ── Rapport Word (si anomalies) ───────────────────────────────

def generate_word_report(rapport: dict, log_file: Path) -> None:
    """Genere un rapport Word si des anomalies ou avertissements sont detectes."""
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        log.warning("python-docx non installe - rapport Word ignore")
        return

    doc = DocxDocument()

    # Titre
    titre = doc.add_heading("OSA Observatory -- Rapport Healthcheck", 0)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ts = rapport["timestamp"][:19]
    statut = rapport["statut_global"]
    doc.add_paragraph(f"Date : {ts}   |   Statut global : {statut}")
    doc.add_paragraph("")

    couleurs_statut = {
        "CRITIQUE": RGBColor(0xC0, 0x00, 0x00),
        "ATTENTION": RGBColor(0xFF, 0x99, 0x00),
        "OK":        RGBColor(0x37, 0x56, 0x23),
    }

    def add_section(title, items, color):
        if not items:
            return
        h = doc.add_heading(title, level=1)
        for run in h.runs:
            run.font.color.rgb = color
        for item in items:
            if isinstance(item, dict):
                p = doc.add_paragraph()
                p.add_run(f"[{item['code']}] ").bold = True
                p.add_run(item["message"])
                p2 = doc.add_paragraph(f"  Action recommandee : {item['action']}")
                p2.runs[0].italic = True
            else:
                doc.add_paragraph(f"  {item}")

    add_section(
        f"Anomalies critiques ({len(rapport['anomalies'])})",
        rapport["anomalies"],
        couleurs_statut["CRITIQUE"]
    )
    add_section(
        f"Avertissements ({len(rapport['avertissements'])})",
        rapport["avertissements"],
        couleurs_statut["ATTENTION"]
    )
    add_section(
        f"Checks OK ({len(rapport['ok'])})",
        rapport["ok"],
        couleurs_statut["OK"]
    )

    # Distribution AMAR 2024
    d = rapport["data"]
    if 2024 in d["amar_dist"]:
        doc.add_heading("Distribution AMAR 2024", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Niveau"
        hdr[1].text = "Nb pays"
        hdr[2].text = "Score moyen"
        for band, vals in sorted(d["amar_dist"][2024].items()):
            row = table.add_row().cells
            row[0].text = band
            row[1].text = str(vals["nb"])
            row[2].text = f"{vals['score']:.3f}"

    # Scores ISA
    doc.add_heading("Scores ISA observes 2020-2024", level=1)
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = "Table Grid"
    h2 = table2.rows[0].cells
    h2[0].text = "Annee"
    h2[1].text = "Score moyen ISA"
    h2[2].text = "Nb pays"
    for y in range(2020, 2025):
        if y in d["isa_scores"]:
            s = d["isa_scores"][y]
            row = table2.add_row().cells
            row[0].text = str(y)
            row[1].text = f"{s['avg']:.5f}"
            row[2].text = str(s["nb"])

    docx_file = REPORT_DIR / f"healthcheck_{ts[:10].replace('-','')}.docx"
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_file))
    log.info("Rapport Word genere : %s", docx_file)
    print(f"  Rapport Word : {docx_file}")


# ── Main ──────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(
        description="OSA Healthcheck -- Diagnostic automatise pipeline + AMAR + doctrine"
    )
    parser.add_argument("--log", type=str, default=None,
        help="Chemin du fichier log SQL (defaut : logs/healthcheck/healthcheck_YYYYMMDD_HHMMSS.log)")
    parser.add_argument("--report", action="store_true",
        help="Genere un rapport Word si anomalies ou avertissements detectes")
    parser.add_argument("--json-only", action="store_true",
        help="Sauvegarde uniquement le JSON, pas de console")
    args = parser.parse_args()

    # Fichier log
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    log_file = Path(args.log) if args.log else LOG_DIR / f"healthcheck_{ts}.log"

    # Execution SQL
    try:
        lines = run_healthcheck_sql(log_file)
    except RuntimeError as e:
        log.error("Healthcheck ECHEC : %s", e)
        sys.exit(1)

    # Parse + analyse
    data    = parse_lines(lines)
    rapport = analyse(data)

    # Sauvegarde JSON
    save_json(rapport, log_file)

    # Rapport console
    if not args.json_only:
        print_rapport(rapport)

    # Rapport Word si demande et anomalies presentes
    if args.report and (rapport["anomalies"] or rapport["avertissements"]):
        generate_word_report(rapport, log_file)

    # Code de retour
    if rapport["statut_global"] == "CRITIQUE":
        log.error("Statut CRITIQUE -- %d anomalie(s) detectee(s)", len(rapport["anomalies"]))
        sys.exit(2)
    elif rapport["statut_global"] == "ATTENTION":
        log.warning("Statut ATTENTION -- %d avertissement(s)", len(rapport["avertissements"]))
        sys.exit(1)
    else:
        log.info("Statut OK -- pipeline sain")
        sys.exit(0)


if __name__ == "__main__":
    main()
